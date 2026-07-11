from io import BytesIO

from docx import Document as DocxDocument

from app.modules.documents.engine.builder_document import BuilderDocument, BuilderSection


def parse_docx(content: bytes) -> BuilderDocument:
    docx = DocxDocument(BytesIO(content))
    sections: list[BuilderSection] = []
    current_title = "Преамбула"
    current_lines: list[str] = []
    section_types = {
        "шапка": "header",
        "предмет": "subject",
        "стоимость": "price",
        "ответственность": "responsibility",
        "реквизиты": "requisites",
        "подписи": "signatures",
    }

    def flush_section() -> None:
        nonlocal current_title, current_lines
        if not current_lines:
            return
        lowered = current_title.lower()
        section_type = "generic"
        for marker, mapped in section_types.items():
            if marker in lowered:
                section_type = mapped
                break
        sections.append(
            BuilderSection(
                title=current_title,
                section_type=section_type,
                content="\n".join(current_lines).strip(),
            )
        )
        current_title = "Преамбула"
        current_lines = []

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading") or text.isupper() and len(text) < 120:
            flush_section()
            current_title = text
            continue
        current_lines.append(text)

    flush_section()

    if not sections:
        sections.append(BuilderSection(title="Содержание", section_type="generic", content=""))

    return BuilderDocument(
        title=sections[0].title if sections else "Импортированный документ",
        doc_type="imported",
        sections=sections,
    )
