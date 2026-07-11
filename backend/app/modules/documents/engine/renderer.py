from io import BytesIO

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.modules.documents.engine.builder_document import BuilderDocument
from app.modules.documents.engine.variables import render_document


def export_docx(document: BuilderDocument) -> bytes:
    rendered = render_document(document)
    docx = DocxDocument()
    docx.add_heading(rendered.title, level=0)
    for section in rendered.sections:
        docx.add_heading(section.title, level=1)
        for line in section.content.splitlines():
            docx.add_paragraph(line)
    buffer = BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


def export_pdf(document: BuilderDocument) -> bytes:
    rendered = render_document(document)
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(40, y, rendered.title[:90])
    y -= 30
    pdf.setFont("Helvetica", 11)
    for section in rendered.sections:
        if y < 80:
            pdf.showPage()
            y = height - 50
            pdf.setFont("Helvetica", 11)
        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(40, y, section.title[:90])
        y -= 18
        pdf.setFont("Helvetica", 11)
        for line in section.content.splitlines():
            if y < 60:
                pdf.showPage()
                y = height - 50
                pdf.setFont("Helvetica", 11)
            pdf.drawString(40, y, line[:110])
            y -= 14
        y -= 8
    pdf.save()
    return buffer.getvalue()


def export_html(document: BuilderDocument) -> str:
    rendered = render_document(document)
    parts = [f"<html><head><meta charset='utf-8'><title>{rendered.title}</title></head><body>"]
    parts.append(f"<h1>{rendered.title}</h1>")
    for section in rendered.sections:
        parts.append(f"<h2>{section.title}</h2>")
        for line in section.content.splitlines():
            parts.append(f"<p>{line}</p>")
    parts.append("</body></html>")
    return "".join(parts)
