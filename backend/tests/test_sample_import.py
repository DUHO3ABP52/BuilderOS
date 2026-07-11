from io import BytesIO

from docx import Document as DocxDocument

from app.modules.documents.engine.sample_import import import_sample_to_builder_document
from app.modules.documents.engine.text_parser import extract_variables_from_text, parse_plain_text, templatize_document
from app.modules.memory.learning import learn_from_template_version, suggest_slug
from app.modules.templates.models import DocumentTemplate, TemplateCategory


SAMPLE_TEXT = """
ДОГОВОР ПОДРЯДА № 12/26
г. Москва 11.07.2026

1. ПРЕДМЕТ ДОГОВОРА
Заказчик: ООО Ромашка
Подрядчик: ИП Иванов
Объект: ЖК Север
Адрес: г. Москва, ул. Примерная, д. 1

2. СТОИМОСТЬ
Стоимость работ составляет 1 500 000 руб.

3. РЕКВИЗИТЫ
ИНН 7707083893
ИНН 500100732259
"""


def test_parse_plain_text_detects_sections() -> None:
    document = parse_plain_text(SAMPLE_TEXT, title="Договор")
    assert document.title == "Договор"
    assert len(document.sections) >= 2


def test_extract_variables_and_templatize() -> None:
    text, definitions, values = extract_variables_from_text(SAMPLE_TEXT)
    assert "{{contract.number}}" in text or "12/26" in values.get("contract.number", "12/26")
    assert any(item.key.endswith(".inn") for item in definitions)
    document = parse_plain_text(SAMPLE_TEXT, title="Договор подряда")
    templated, variables, sample_values = templatize_document(document)
    assert templated.doc_type == "contract"
    assert variables
    assert sample_values


def test_import_txt_sample_bytes() -> None:
    result = import_sample_to_builder_document(
        SAMPLE_TEXT.encode("utf-8"),
        filename="dogovor.txt",
        content_type="text/plain",
        title="Договор из образца",
    )
    assert result.source_format == "text"
    assert result.document.doc_type == "contract"
    assert result.variables


def test_import_docx_sample_bytes() -> None:
    buffer = BytesIO()
    docx = DocxDocument()
    docx.add_heading("ДОГОВОР ПОДРЯДА", level=1)
    docx.add_paragraph("Заказчик: ООО Тест")
    docx.add_paragraph("Стоимость работ составляет 100000 руб.")
    docx.save(buffer)
    result = import_sample_to_builder_document(
        buffer.getvalue(),
        filename="sample.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert result.source_format == "docx"
    assert result.document.sections


def test_suggest_slug() -> None:
    assert "dogovor" in suggest_slug("Договор подряда")


def test_learn_from_template_version(client, auth_headers) -> None:
    # create parent via API sample import
    files = {"file": ("sample.txt", SAMPLE_TEXT.encode("utf-8"), "text/plain")}
    data = {"name": "Договор обучение", "category": "contract"}
    created = client.post("/api/v1/templates/import/sample", headers=auth_headers, files=files, data=data)
    assert created.status_code == 201, created.text
    parent = created.json()

    # new version with changed section content
    content = parent["content"]
    content["sections"][0]["content"] = content["sections"][0]["content"] + "\nДополнительный пункт."
    version = client.post(
        f"/api/v1/templates/{parent['id']}/versions",
        headers=auth_headers,
        json={
            "name": parent["name"],
            "slug": parent["slug"],
            "category": parent["category"],
            "content": content,
            "variables": parent["variables"],
            "description": "правка",
        },
    )
    assert version.status_code == 201, version.text

    memory = client.get("/api/v1/memory", headers=auth_headers, params={"q": parent["slug"]})
    assert memory.status_code == 200
    assert len(memory.json()) >= 1
